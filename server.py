import os
import json
import tempfile
import uuid
from datetime import date
from io import BytesIO
from pathlib import Path

from flask import Flask, request, jsonify, send_file, send_from_directory
import anthropic
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

# ─── CONFIGURACAO ─────────────────────────────────────────────────────────────
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
GROQ_API_KEY      = os.environ.get("GROQ_API_KEY", "")
# ──────────────────────────────────────────────────────────────────────────────

client      = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
groq_client = Groq(api_key=GROQ_API_KEY)

AZUL   = RGBColor(0x1A, 0x3A, 0x5C)
LARANJ = RGBColor(0xE8, 0x77, 0x22)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA  = RGBColor(0x88, 0x88, 0x88)

PASTA_SAIDA = Path(tempfile.gettempdir()) / "agiledog_output"
PASTA_SAIDA.mkdir(exist_ok=True)

app = Flask(__name__, static_folder="static", static_url_path="")


# ─── HELPERS DOCX (identicos ao app.py original) ─────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(paragraph, text, bold=False, italic=False,
            size=11, color=None, font="Arial"):
    run = paragraph.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_border_bottom(paragraph, hex_color, size=12):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_space(doc, size_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    add_run(p, "", size=size_pt)

def add_highlight_box(doc, texto):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "D6E4F0")
    set_cell_borders(cell, color="A8C4E0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, texto, size=11)
    add_space(doc, 6)

def add_step_table(doc, itens):
    if not itens:
        return
    table = doc.add_table(rows=len(itens), cols=2)
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Largura total da tabela
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "8880")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # Layout fixo: garante que as larguras de coluna sejam respeitadas
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    for i, item in enumerate(itens):
        # Coluna do numero: estreita (600 dxa = ~1.06 cm)
        cell_n = table.cell(i, 0)
        set_cell_bg(cell_n, "E87722")
        set_cell_borders(cell_n, color="E87722")
        tcPr_n = cell_n._tc.get_or_add_tcPr()
        tcW_n = OxmlElement("w:tcW")
        tcW_n.set(qn("w:w"), "600")
        tcW_n.set(qn("w:type"), "dxa")
        tcPr_n.append(tcW_n)
        p_n = cell_n.paragraphs[0]
        p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_n.paragraph_format.space_before = Pt(6)
        p_n.paragraph_format.space_after  = Pt(6)
        add_run(p_n, str(i + 1), bold=True, size=12, color=BRANCO)
        # Coluna do texto: restante
        cell_t = table.cell(i, 1)
        tcPr_t = cell_t._tc.get_or_add_tcPr()
        tcW_t = OxmlElement("w:tcW")
        tcW_t.set(qn("w:w"), "8280")
        tcW_t.set(qn("w:type"), "dxa")
        tcPr_t.append(tcW_t)
        set_cell_bg(cell_t, "F2F2F2")
        set_cell_borders(cell_t, color="CCCCCC")
        p_t = cell_t.paragraphs[0]
        p_t.paragraph_format.space_before = Pt(5)
        p_t.paragraph_format.space_after  = Pt(5)
        p_t.paragraph_format.left_indent  = Pt(4)
        add_run(p_t, item, size=11)
    add_space(doc, 6)

def add_section_title(doc, titulo, numero=None):
    add_space(doc, 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    add_border_bottom(p, "E87722", size=8)
    label = f"{numero}. {titulo}" if numero else titulo
    add_run(p, label.upper(), bold=True, size=12, color=LARANJ)

def add_bullet_list(doc, itens):
    for item in itens:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_run(p, item, size=11)
    add_space(doc, 4)

def add_text_block(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, texto, size=11)

def build_header_footer(doc):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    cab = header.paragraphs[0]
    cab.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_border_bottom(cab, "1A3A5C", size=18)
    add_run(cab, "AGILE DOG", bold=True, size=14, color=AZUL)
    add_run(cab, "   |   ", size=10, color=CINZA)
    r = cab.add_run("Protocolo de treinamento")
    r.italic         = True
    r.font.name      = "Arial"
    r.font.size      = Pt(10)
    r.font.color.rgb = CINZA
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p.clear()
    rod = footer.paragraphs[0]
    add_border_bottom(rod, "E87722", size=6)
    rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(rod, "Agile Dog  •  Desenvolvido por Kaique Rocha  •  v1.0", size=9, color=CINZA)


# ─── NOVA FUNCAO: INSERCAO DE FOTO NO DOCUMENTO ──────────────────────────────

def inserir_foto_protocolo(doc, foto_bytes):
    """
    Insere foto do animal: recorte quadrado centralizado, moldura elegante azul.
    """
    img = PILImage.open(BytesIO(foto_bytes))

    # Converter para RGB
    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")

    # Corrigir orientacao EXIF (fotos de celular)
    try:
        import PIL.ExifTags
        exif = img._getexif()
        if exif:
            for tag, val in exif.items():
                if PIL.ExifTags.TAGS.get(tag) == "Orientation":
                    rotacoes = {3: 180, 6: 270, 8: 90}
                    if val in rotacoes:
                        img = img.rotate(rotacoes[val], expand=True)
                    break
    except Exception:
        pass

    # Recorte quadrado centralizado
    w, h = img.size
    lado = min(w, h)
    esq  = (w - lado) // 2
    top  = (h - lado) // 2
    img  = img.crop((esq, top, esq + lado, top + lado))

    # Redimensionar para 600x600
    img = img.resize((600, 600), PILImage.LANCZOS)

    # Salvar em temp
    tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
    img.save(tmp.name, "JPEG", quality=90)
    tmp.close()

    # Moldura elegante: tabela 3x3 simulando padding + borda dupla
    add_space(doc, 6)

    # Tabela externa (borda laranja fina)
    t_ext = doc.add_table(rows=1, cols=1)
    t_ext.style = "Table Grid"
    c_ext = t_ext.cell(0, 0)
    set_cell_bg(c_ext, "E87722")
    set_cell_borders(c_ext, color="E87722")
    tcPr_e = c_ext._tc.get_or_add_tcPr()
    tcW_e = OxmlElement("w:tcW")
    tcW_e.set(qn("w:w"), "4680")
    tcW_e.set(qn("w:type"), "dxa")
    tcPr_e.append(tcW_e)

    # Paragrafo interno com a imagem
    p_ext = c_ext.paragraphs[0]
    p_ext.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ext.paragraph_format.space_before = Pt(4)
    p_ext.paragraph_format.space_after  = Pt(4)

    # Tabela interna (fundo branco, moldura azul escuro)
    t_int = doc.add_table(rows=1, cols=1)
    t_int.style = "Table Grid"
    c_int = t_int.cell(0, 0)
    set_cell_bg(c_int, "FFFFFF")

    # Borda azul mais grossa
    tc_int    = c_int._tc
    tcPr_int  = tc_int.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "18")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "1A3A5C")
        tcBorders.append(el)
    tcPr_int.append(tcBorders)

    tcW_i = OxmlElement("w:tcW")
    tcW_i.set(qn("w:w"), "4536")
    tcW_i.set(qn("w:type"), "dxa")
    tcPr_int.append(tcW_i)

    p_int = c_int.paragraphs[0]
    p_int.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_int.paragraph_format.space_before = Pt(0)
    p_int.paragraph_format.space_after  = Pt(0)
    run = p_int.add_run()
    run.add_picture(tmp.name, width=Cm(8))

    os.unlink(tmp.name)

    # Centralizar as duas tabelas
    for t in [t_ext, t_int]:
        tbl   = t._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        tblPr.append(jc)

    add_space(doc, 8)


# ─── TRANSCRICAO (identica ao app.py original) ───────────────────────────────

def transcrever_audio(caminho_audio):
    with open(caminho_audio, "rb") as f:
        result = groq_client.audio.transcriptions.create(
            file=(Path(caminho_audio).name, f),
            model="whisper-large-v3",
            language="pt",
            response_format="text"
        )
    return result


# ─── GERACAO DE CONTEUDO JSON (identica ao app.py original) ──────────────────

def gerar_protocolo_json(transcript):
    system = """Voce e um adestrador canino especialista na metodologia Equilibrio dos Reforcamentos (Agile Dog).
Recebe a transcricao de uma aula pratica ao vivo.

INSTRUCOES:
- Extraia o nome do cao e do tutor da transcricao. Procure por frases como "o [nome]", "seu [nome]", ou o tutor sendo chamado pelo nome.
- IGNORE conversas informais e trechos sem relacao com o treino
- Crie secoes: Objetivo da Proxima Sessao, Exercicios Recomendados, Criterios de Reforcamento, Observacoes Tecnicas
- Para Exercicios Recomendados use tipo "passos" com cada passo descrito claramente
- Para outras secoes use tipo "texto" ou "lista"
- Adicione 1 a 2 "destaques" com dicas importantes para o tutor
- Linguagem clara, sem jargao tecnico, tom profissional

Retorne SOMENTE este JSON valido, sem markdown, sem texto antes ou depois:
{
  "nome_cao": "string",
  "nome_tutor": "string",
  "subtitulo": "string",
  "intro": "string",
  "secoes": [
    {"titulo": "string", "tipo": "texto|passos|lista", "conteudo": "string", "itens": ["string"]}
  ],
  "destaques": ["string"],
  "fechamento": "string"
}"""

    resposta = client.messages.create(
        model="claude-sonnet-4-5", max_tokens=3000, system=system,
        messages=[{"role": "user", "content": f"Transcricao:\n\n{transcript}"}]
    )
    raw = resposta.content[0].text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    raw = raw[raw.find("{"):raw.rfind("}")+1]
    return json.loads(raw)


def gerar_relatorio_json(transcript):
    system = """Voce e um assistente especializado em adestramento canino positivo (Agile Dog).
Recebe a transcricao de uma aula pratica ao vivo.

INSTRUCOES:
- Extraia o nome do cao e do tutor da transcricao
- IGNORE comentarios casuais e conversas paralelas
- Crie secoes: Resumo da Aula, Comportamentos Trabalhados, Evolucao Observada, Pontos de Atencao, Proximos Passos
- Use tipo "texto" para resumos e "lista" para itens enumeraveis
- Linguagem acessivel, tom encorajador

Retorne SOMENTE este JSON valido, sem markdown, sem texto antes ou depois:
{
  "nome_cao": "string",
  "nome_tutor": "string",
  "subtitulo": "string",
  "secoes": [
    {"titulo": "string", "tipo": "texto|lista", "conteudo": "string", "itens": ["string"]}
  ],
  "fechamento": "string"
}"""

    resposta = client.messages.create(
        model="claude-sonnet-4-5", max_tokens=3000, system=system,
        messages=[{"role": "user", "content": f"Transcricao:\n\n{transcript}"}]
    )
    raw = resposta.content[0].text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    raw = raw[raw.find("{"):raw.rfind("}")+1]
    return json.loads(raw)


# ─── GERACAO DOCX (salvar_protocolo expandido com foto opcional) ──────────────

def salvar_protocolo(dados, nome_arquivo, foto_bytes=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    build_header_footer(doc)

    nome_cao   = dados.get("nome_cao",   "Nao identificado")
    nome_tutor = dados.get("nome_tutor", "Nao identificado")

    add_space(doc, 6)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(2)
    add_run(p_titulo, "Protocolo de Treinamento", bold=True, size=18, color=AZUL)

    subtitulo = dados.get("subtitulo", "")
    if subtitulo:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(6)
        add_run(p_sub, subtitulo, bold=True, size=13, color=AZUL)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(10)
    add_run(p_info, f"{nome_cao}  |  {nome_tutor}  |  {date.today().strftime('%d/%m/%Y')}",
            size=10, color=CINZA)

    # ── FOTO: inserida na area de identificacao, apenas se disponivel ──
    if foto_bytes:
        inserir_foto_protocolo(doc, foto_bytes)

    intro = dados.get("intro", "")
    if intro:
        p_intro = doc.add_paragraph()
        p_intro.paragraph_format.space_after = Pt(8)
        r = p_intro.add_run(intro)
        r.italic      = True
        r.font.name   = "Arial"
        r.font.size   = Pt(11)

    destaques = dados.get("destaques", [])
    secoes    = dados.get("secoes", [])

    for i, secao in enumerate(secoes, 1):
        add_section_title(doc, secao.get("titulo", ""), numero=i)
        tipo = secao.get("tipo", "texto")
        if tipo == "passos":
            add_step_table(doc, secao.get("itens", []))
        elif tipo == "lista":
            add_bullet_list(doc, secao.get("itens", []))
        else:
            add_text_block(doc, secao.get("conteudo", ""))
        if i <= len(destaques):
            add_highlight_box(doc, destaques[i - 1])

    for d in destaques[len(secoes):]:
        add_highlight_box(doc, d)

    fechamento = dados.get("fechamento", "")
    if fechamento:
        add_space(doc, 8)
        p = doc.add_paragraph()
        add_run(p, fechamento, bold=True, size=11)

    doc.save(nome_arquivo)


def salvar_relatorio(dados, nome_arquivo):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    build_header_footer(doc)

    nome_cao   = dados.get("nome_cao",   "Nao identificado")
    nome_tutor = dados.get("nome_tutor", "Nao identificado")

    add_space(doc, 6)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(2)
    add_run(p_titulo, "Relatorio de Desempenho", bold=True, size=18, color=AZUL)

    subtitulo = dados.get("subtitulo", "")
    if subtitulo:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(6)
        add_run(p_sub, subtitulo, bold=True, size=13, color=AZUL)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(10)
    add_run(p_info, f"{nome_cao}  |  {nome_tutor}  |  {date.today().strftime('%d/%m/%Y')}",
            size=10, color=CINZA)

    for i, secao in enumerate(dados.get("secoes", []), 1):
        add_section_title(doc, secao.get("titulo", ""), numero=i)
        tipo = secao.get("tipo", "texto")
        if tipo == "lista":
            add_bullet_list(doc, secao.get("itens", []))
        else:
            add_text_block(doc, secao.get("conteudo", ""))

    fechamento = dados.get("fechamento", "")
    if fechamento:
        add_space(doc, 8)
        p = doc.add_paragraph()
        add_run(p, fechamento, bold=True, size=11)

    doc.save(nome_arquivo)


# ─── ROTAS FLASK ─────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory("static", "index.html")

@app.route("/manifest.json")
def manifest():
    return send_from_directory("static", "manifest.json")

@app.route("/sw.js")
def service_worker():
    return send_from_directory("static", "sw.js",
                               mimetype="application/javascript")


@app.route("/api/processar", methods=["POST"])
def processar():
    """
    Recebe:
      - audio: arquivo de audio (multipart/form-data)
      - foto:  arquivo de imagem opcional (multipart/form-data)
    Retorna JSON com IDs dos documentos gerados.
    """
    if "audio" not in request.files:
        return jsonify({"erro": "Arquivo de audio nao enviado."}), 400

    audio_file = request.files["audio"]
    foto_file  = request.files.get("foto")

    # Salvar audio em temp
    sufixo_audio = Path(audio_file.filename).suffix or ".mp4"
    tmp_audio = tempfile.NamedTemporaryFile(suffix=sufixo_audio, delete=False)
    audio_file.save(tmp_audio.name)
    tmp_audio.close()

    # Ler foto se enviada
    foto_bytes = None
    if foto_file and foto_file.filename:
        foto_bytes = foto_file.read()

    try:
        # 1. Transcricao
        transcript = transcrever_audio(tmp_audio.name)

        # 2. Gerar JSONs via Claude
        dados_prot = gerar_protocolo_json(transcript)
        dados_rel  = gerar_relatorio_json(transcript)

        nome_cao   = dados_prot.get("nome_cao",   "Cao")
        nome_tutor = dados_prot.get("nome_tutor", "Tutor")

        # 3. Nomes de arquivo
        data_hoje = date.today().strftime("%Y%m%d")
        nome_base = nome_cao.replace(" ", "_")
        sessao_id = str(uuid.uuid4())[:8]

        path_prot = PASTA_SAIDA / f"protocolo_{nome_base}_{data_hoje}_{sessao_id}.docx"
        path_rel  = PASTA_SAIDA / f"relatorio_{nome_base}_{data_hoje}_{sessao_id}.docx"

        # 4. Gerar documentos
        salvar_protocolo(dados_prot, str(path_prot), foto_bytes=foto_bytes)
        salvar_relatorio(dados_rel,  str(path_rel))

        return jsonify({
            "ok":        True,
            "nome_cao":  nome_cao,
            "nome_tutor": nome_tutor,
            "protocolo": path_prot.name,
            "relatorio": path_rel.name,
        })

    except Exception as e:
        import traceback
        return jsonify({"erro": str(e), "detalhe": traceback.format_exc()}), 500

    finally:
        os.unlink(tmp_audio.name)


@app.route("/api/download/<nome_arquivo>")
def download(nome_arquivo):
    """Serve o arquivo gerado para download."""
    caminho = PASTA_SAIDA / nome_arquivo
    if not caminho.exists():
        return jsonify({"erro": "Arquivo nao encontrado."}), 404
    return send_file(str(caminho), as_attachment=True,
                     download_name=nome_arquivo,
                     mimetype="application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
